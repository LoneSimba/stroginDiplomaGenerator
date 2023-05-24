import json
import os
import sys

import unicodedata

import openpyxl
import re
from docx import Document as D
from docx.document import Document
from docx.table import Table
from openpyxl import Workbook
from openpyxl.worksheet.worksheet import Worksheet

from src import MailCloudDownloader


# main() start
def main():
    print('Loading config and ingest file\n')

    config: dict = load_config()
    comp: str = config.get("comp")
    ingest_path: str = os.path.join(os.getcwd(), "files/", comp)

    # table_name = MailCloudDownloader.download_table(config.get('table_link'), ingest_path)
    table_name = config.get('table_name')

    if table_name is None:
        raise Exception('No table was downloaded')

    print('Load complete, ingest table found: ' + table_name + '\n')

    table: Workbook = openpyxl.load_workbook(os.path.join(ingest_path, table_name))
    table_conf: dict = config.get('table').get(comp)

    print('Processing nominations...\n')

    for i in range(0, table_conf.get('cats')):
        ws: Worksheet = table.worksheets[i]
        # if not re.findall(r'\d\s', str(ws.title)):
        #     continue

        title: str = re.sub(r'\d\s', '', str(ws.title))
        print('[' + str(i+1) + '/' + str(table_conf.get('cats')) + '] Processing nomination: ' + title)

        start_row: int = 5
        processed: int = 0

        for row in ws.iter_rows(min_row=start_row, values_only=True):
            processed += 1

            if not row[0]:
                break

            values: list = [row[v] for k, v in table_conf.get('cols').items()]
            entry: dict = dict(zip(table_conf.get('cols').keys(), values))

            # if any(res in entry.get('result') for res in ['Специальный приз', 'Лауреат I степени', 'Гран-при']):
            if any(res in entry.get('result') for res in ['Лауреат II', 'Дипломант', 'Участник']):

                # clean garbage from names
                participant = re.sub(r'\"|\'|\«|\»', '', str(entry.get('participant')))
                tutor = re.sub(r'\"|\'|\«|\»', '', str(entry.get('tutor')))
                entry.update({'tutor': tutor, 'participant': participant})

                count: int = int(re.findall(r'\d+', str(entry.get('count')))[0])
                if count < 1:
                    continue

                nomination = re.sub(r'\"|\«|\»', '', entry.get('nomination'))
                title = re.sub(r'^\"|\"$|\"\.$|^\«|\»$|\».$|^\'|\'$|\'.$', '', str(entry.get("title")))
                entry.update({'nomination': nomination, "title": title})

                is_group: bool = count > 1
                is_individual: bool = 'Самостоятельный участник' in (entry.get('form') or '')
                is_prized: bool = any(x in entry.get('result') for x in ['Дипломант', 'Лауреат', 'Гран-при', 'Специальный'])
                is_sp: bool = 'Специальный приз' in entry.get('result')
                is_gp: bool = 'Гран-при' in entry.get('result')

                output_path: str = os.path.join(
                    # os.getcwd(), "output", comp, str(ws.title),
                    os.getcwd(), "output", comp,
                    str((entry.get('tutor'), '')[is_individual]).strip()
                )

                in_filename: str = "шаблон_"
                in_filename += ("благодарность", "диплом")[is_prized]

                if is_sp:
                    in_filename += "_спец"
                elif is_gp:
                    in_filename += "_гранпри"
                else:
                    in_filename += "_участник"

                in_filename += ("", "_группа")[is_group and not is_individual]
                in_filename += ("", "_инд")[is_individual]
                in_filename += ".docx"

                if not is_individual:
                    tutor_ent: dict = {'tutor': entry.get('tutor') or '', 'group': str(entry.get('group')) or '',
                                       'school': str(entry.get('school')) or ''}
                    create_tutor_dipl(tutor_ent, ingest_path, output_path, is_gp, is_sp)

                # if not is_group:
                create_diploma(entry, os.path.join(ingest_path, in_filename), output_path, is_prized)
                    # continue

                # participants = re.sub(r',', ', ', str(entry.get('participant')))
                # participants = participants.split(',')

                # if len(participants) != count:
                    # create_diploma(entry, os.path.join(ingest_path, in_filename), output_path, is_prized)
                    # continue

                # for partic in participants:
                #     en = dict(entry)
                #     en.update({'participant': partic})
                #     create_diploma(en, os.path.join(ingest_path, in_filename), output_path, is_prized)

        print("Nomination done, " + str(processed) + " processed\n")

# main() end


def create_diploma(entry: dict, ingest_path: str, output_path: str, is_prized: bool):

    part: str = re.sub(r'[\"\'\\/?%*:|<>\n]', '', str(entry.get('participant')))

    if len(part) > 64:
        part = part[:64]

    tit: str = re.sub(r'[\"\'\\/?%*:|<>\n]', '', str(entry.get('title')))

    if len(tit) > 64:
        tit = tit[:64]

    out_filename: str = ("благодарность", "диплом")[is_prized]
    out_filename += re.sub(r'[\"\'\\/?%*:|<>\n]', '',  (" " + part + " " + tit))

    output_path = unicodedata.normalize('NFKD', output_path)
    out_filename = unicodedata.normalize('NFKD', os.path.join(output_path, out_filename + ".docx"))

    temp: Document = D(ingest_path)
    keys: list = list(entry.keys())
    table: Table = temp.tables[0]
    for cell in table.column_cells(0):
        par = cell.paragraphs[0]
        for run in par.runs:
            text = str(run.text).replace('%', '')
            if text in keys:
                run.text = str(entry.get(text)).replace('None', '') or ''

    os.makedirs(output_path, exist_ok=True)
    temp.save(out_filename)


def create_tutor_dipl(entry: dict, ingest_path: str, output_path: str, is_gp: bool, is_sp: bool):
    out_filename: str = "благодарность"
    out_filename += re.sub(r'["\'\\/?%*:|<>\n]', '',
                           (" " + str(entry.get('tutor')) + " " + entry.get('group') + " " + entry.get('school')))
    out_filename += ("", " гран при")[is_gp]
    out_filename += ("", " спец приз")[is_sp]
    out_filename += ".docx"
    output_path = unicodedata.normalize('NFKD', output_path)
    out_filename = unicodedata.normalize('NFKD', os.path.join(output_path, out_filename))

    if os.path.isfile(out_filename):
        return

    temp_name: str = "шаблон_благодарность"

    if is_sp:
        temp_name += "_спец"
    elif is_gp:
        temp_name += "_гранпри"
    else:
        temp_name += "_участник"

    temp_name += "_педагог.docx"
    temp: Document = D(os.path.join(ingest_path, temp_name))
    keys: list = list(entry.keys())
    table: Table = temp.tables[0]
    for c in table.column_cells(0):
        p = c.paragraphs[0]
        for r in p.runs:
            text = str(r.text).replace('%', '')
            if text in keys:
                r.text = str(entry.get(text)).replace('None', '') or ''
    os.makedirs(output_path, exist_ok=True)
    temp.save(os.path.join(output_path, out_filename))


# load_config() start
def load_config() -> dict:
    try:
        with open('conf.json', 'r', encoding='utf-8') as conf:
            config: dict = json.load(conf)
    except OSError:
        print("Failed to open conf.json file")
        exit(1)
    except json.JSONDecodeError:
        print("conf.json contents are broken")
        exit(1)

    return config


# load_config() end


if __name__ == '__main__':
    # try:
        main()
        print("Done")
    # except Exception as e:
    #     exc_type, exc_obj, exc_tb = sys.exc_info()
    #     f = exc_tb.tb_frame
    #     file = f.f_code.co_filename
    #
    #     print("Fatal error: " + str(e) + " in " + file + " at " + str(exc_tb.tb_lineno))
